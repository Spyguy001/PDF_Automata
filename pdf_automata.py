""" Program: PDF Automata by Amrital Aujla
    This program takes a Driver Settlement issued to drivers for Falcon Motor Freight Ltd. and turns it into a
    pay report that the drivers use to pay others.
    
    To use it, make two folders in the folder this program is in, one called pdf_storage and the other called
    doc_storage. Put the Driver Settlement, renamed as 'falcon' into the pdf_storage folder. Also put in there a text
    file called rate that has the driver's rate written in it only. Run the program and the produced file will be
    in the doc_storage folder.
"""

import PyPDF2
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import docx
from docx.shared import Cm


def main():
    """ Takes a Driver Settlement PDF and makes a Pay Report docx """
    pdf_obj = open('pdf_storage/falcon.pdf', 'rb')
    pdf = PyPDF2.PdfFileReader(pdf_obj)
    all_text = get_full_text(pdf)

    if all_text == '':
        doc = docx.Document()
        doc.add_paragraph('Nothing Found in PDF')
        doc.save('doc_storage/Error.docx')
    else:
        text = get_relevant_text(all_text)
        name_and_date = get_name_date_data(text)
        rate = get_rate()
        trip_list = make_trip_list(text)
        pickups_and_miles = get_trip_blocks(text, pickups_block=True)
        pay_calculations = calculate_pay(pickups_and_miles, rate)
        make_document(name_and_date, trip_list, rate, pickups_and_miles, pay_calculations)


def get_full_text(pdf):
    """ Gets the full text from all the pages of the pdf """
    num_pages = pdf.numPages
    count = 0
    all_text = ''

    while count < num_pages:
        page_obj = pdf.getPage(count)
        count += 1
        all_text += page_obj.extractText()

    return all_text


def get_relevant_text(all_text):
    """ Gets the important parts of the text and discards punctuation and certain words """
    tokens = word_tokenize(all_text)
    punctuations = ['(', ')', ';', ':', '[', ']', ',', '-']
    stop_words = stopwords.words('english')
    text = [word for word in tokens if word not in stop_words and word not in punctuations]
    return text


def get_name_date_data(text):
    """ Gets the name and date from the text and returns them as a list """
    name = text[9] + ' ' + text[10]
    date = text[12] + ' ' + text[13] + ' ' + text[14][:4] + ' to ' + text[15] + ' ' + text[16] + ' ' + text[17][:4]
    return [name, date]


def get_rate():
    """ Gets the rate for the trip from a text file """
    file = open('pdf_storage/rate.txt')
    rate = file.readline()
    file.close()
    if rate == '':
        return 24
    return int(rate)


def make_trip_list(text):
    """ Makes a list of each individual trips, detailing from where to where they went """
    trip_blocks = get_trip_blocks(text)
    trip_list = []
    for trip in trip_blocks:
        trip_list.append(get_trip_summary(trip))
    return trip_list


def get_trip_blocks(text, pickups_block=False):
    """ Gets the part of the text list relevant only to the individual trips
        If pickups_block is True, gets the part of the text relevant to the pickups """
    chopped_list = []
    chop_index = -1
    for i in range(0, len(text)):
        if "KM'sTrip" in text[i]:
            chop_index += 1
            chopped_list.append([])
            chopped_list[chop_index].append(text[i])
        elif chop_index >= 0 and text[i] != 'Trip':
            chopped_list[chop_index].append(text[i])
        if text[i] == 'Trip' and text[i + 1] == 'NoDescriptionQtyRateCAD':
            if pickups_block is True:
                return get_pickups_and_miles(text[i:])
            break
    return chopped_list


def get_trip_summary(trip):
    """ Gets the trip summary in the format of 'where to where' from a certain block of trip data """
    truck_number = trip[1][6:10]
    pickup_location = ''
    delivery_location = ''
    exit_counter = False
    for i in range(0, len(trip)):
        if 'PICKUP' in trip[i] and exit_counter is False:
            full_term = trip[i - 1] + ''
            pickup_location = get_location(full_term, i, truck_number, trip)
            exit_counter = True
        if 'DELIVER' in trip[i] and exit_counter is True:
            full_term = trip[i - 1] + ''
            delivery_location = get_location(full_term, i, truck_number, trip)
            break
    return pickup_location + ' to ' + delivery_location


def get_location(full_term, i, truck_number, trip):
    """ Gets the full name of the location for delivery or pickup """
    if truck_number in full_term:
        return full_term[full_term.find(truck_number)+4:] + ', ' + trip[i][:2]
    else:
        partial_term = full_term + ''
        full_term = trip[i - 2]
        if truck_number in full_term:
            return full_term[full_term.find(truck_number)+4:] + ' ' + partial_term + ', ' + trip[i][:2]
        else:
            partial_term = partial_term + full_term
            full_term = trip[i - 3]
            return full_term[full_term.find(truck_number)+4:] + ' ' + partial_term + ', ' + trip[i][:2]


def get_pickups_and_miles(text):
    """ Gets the number of pickups and amount of miles per trip from a shortened list """
    miles_list = []
    pickups_list = []
    for i in range(len(text)):
        if 'PICKUP' in text[i]:
            pickups_list.append(int(float(text[i + 1].replace(',', ''))))
        elif 'MILEAGE' in text[i]:
            if len(pickups_list) <= len(miles_list):
                pickups_list.append(0)
            miles_list.append(int(float(text[i + 1].replace(',', ''))))
        if 'DateSupplier' in text[i]:
            break
    return [pickups_list, miles_list]


def calculate_pay(pickups_and_miles, rate):
    """ Calculates the total pay for the driver and gives back all the relevant info as well
        @type rate: int
        @type pickups_and_miles: list
    """
    total_miles = sum(pickups_and_miles[1])
    total_pickups = sum(pickups_and_miles[0]) * 5
    total_pay = round(total_miles * (rate / 100), 2)
    hst = round(total_pay * 0.13, 2)
    final_pay = round(total_pay + hst + total_pickups)
    return [total_miles, total_pay, total_pickups, hst, final_pay]


def make_document(name_and_date, trip_list, rate, pickups_and_miles, pay_calculations):
    """ Makes a word document arranging all the info gotten from the pdf into a pay report """
    doc = docx.Document()
    doc.add_paragraph(name_and_date[0])
    table = make_trip_table(doc, trip_list)
    add_trip_details(table, name_and_date, trip_list, rate, pickups_and_miles)
    add_pay_details(table, trip_list, pay_calculations)
    doc.save('doc_storage/' + name_and_date[0] + ' ' + name_and_date[1] + '.docx')


def make_trip_table(doc, trip_list):
    """ Makes a table with all the necessary rows and columns and attributes and returns it """
    table = doc.add_table(rows=len(trip_list) + 6, cols=0)
    table.style = doc.styles['TableGrid']
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)
    table.add_column(Cm(5.2))
    table.add_column(Cm(7.35))
    table.add_column(Cm(1.6))
    table.add_column(Cm(2.1))
    table.add_column(Cm(2.1))

    table.cell(0, 0).add_paragraph('').add_run('DATE').bold = True
    table.cell(0, 1).add_paragraph('').add_run('ORDER').bold = True
    table.cell(0, 2).add_paragraph('').add_run('MILES').bold = True
    table.cell(0, 3).add_paragraph('').add_run('RATE (cents)').bold = True
    table.cell(0, 4).add_paragraph('').add_run('PICKUPS').bold = True
    table.cell(-5, 0).add_paragraph('').add_run('TOTAL MILES').bold = True
    table.cell(-4, 0).add_paragraph('').add_run('TOTAL PAY').bold = True
    table.cell(-3, 0).add_paragraph('').add_run('TOTAL PICKUPS').bold = True
    table.cell(-2, 0).add_paragraph('').add_run('HST').bold = True
    table.cell(-1, 0).add_paragraph('').add_run('FINAL PAY').bold = True

    for i in range(1, 6):
        table.cell(0 - i, 1).merge(table.cell(0 - i, 4))
    return table


def add_trip_details(table, name_and_date, trip_list, rate, pickups_and_miles):
    """ Adds the details about the trip and each of their miles and pickups to the table """
    trip_amount = len(trip_list)
    for x in range(0, trip_amount):
        table.cell(x + 1, 0).add_paragraph(name_and_date[1])
        table.cell(x + 1, 1).add_paragraph(trip_list[x])
        table.cell(x + 1, 2).add_paragraph(str(pickups_and_miles[1][x]))
        table.cell(x + 1, 3).add_paragraph(str(rate))
        table.cell(x + 1, 4).add_paragraph(str(pickups_and_miles[0][x]))


def add_pay_details(table, trip_list, pay_calculations):
    """ Adds the details about the pay near the end of the table """
    for i in range(5):
        table.cell(len(trip_list) + i + 1, 1).add_paragraph(str(pay_calculations[i]))


if __name__ == '__main__':
    main()
